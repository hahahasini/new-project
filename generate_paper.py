"""
Generate JCES-format research paper for VitaDetect project.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

IMAGES_DIR = r"d:\All project files\All project files\new project\images"
OUTPUT_PATH = r"d:\All project files\All project files\new project\VitaDetect_JCES_Paper.docx"

doc = Document()

# ── Page Setup ──
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(10)
rPr = style.element.get_or_add_rPr()
rFonts = rPr.makeelement(qn('w:rFonts'), {qn('w:eastAsia'): 'Times New Roman'})
rPr.append(rFonts)

# ── Helper Functions ──
def add_journal_header(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Journal of Computer and Electrical Sciences (JCES) - Volume-1 Issue-3, March 2026")
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'
    run.italic = True
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("ISSN: 3049-2602(Online)")
    run2.font.size = Pt(9)
    run2.font.name = 'Times New Roman'
    run2.italic = True

def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(6)
    p.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

def add_authors(doc, authors_text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(2)
    run = p.add_run(authors_text)
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'

def add_affiliation(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'
    run.italic = True

def add_section_heading(doc, number, title):
    p = doc.add_paragraph()
    p.space_before = Pt(12)
    p.space_after = Pt(6)
    run = p.add_run(f"{number}. {title}")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

def add_body_text(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.space_after = Pt(4)
    pf = p.paragraph_format
    pf.first_line_indent = Cm(1.27)
    pf.line_spacing = 1.15
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.space_after = Pt(2)
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'
    if not p.runs:
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'
    else:
        p.runs[0].text = text

def add_figure(doc, image_path, caption, width=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(6)
    if os.path.exists(image_path):
        run = p.add_run()
        if width:
            run.add_picture(image_path, width=width)
        else:
            run.add_picture(image_path, width=Inches(4.5))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.space_after = Pt(8)
    run = cap.add_run(caption)
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'
    run.bold = True

def add_abstract_block(doc, text, keywords):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.space_before = Pt(6)
    p.space_after = Pt(4)
    pf = p.paragraph_format
    pf.line_spacing = 1.15
    run_label = p.add_run("Abstract: ")
    run_label.bold = True
    run_label.font.size = Pt(10)
    run_label.font.name = 'Times New Roman'
    run_body = p.add_run(text)
    run_body.font.size = Pt(10)
    run_body.font.name = 'Times New Roman'

    pk = doc.add_paragraph()
    pk.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pk.space_after = Pt(8)
    rk_label = pk.add_run("Keywords: ")
    rk_label.bold = True
    rk_label.font.size = Pt(10)
    rk_label.font.name = 'Times New Roman'
    rk_body = pk.add_run(keywords)
    rk_body.font.size = Pt(10)
    rk_body.font.name = 'Times New Roman'
    rk_body.italic = True


# ══════════════════════════════════════════════════════════════════
# BUILD THE PAPER
# ══════════════════════════════════════════════════════════════════

add_journal_header(doc)

add_title(doc, "VITADETECT: AN AI-POWERED VITAMIN DEFICIENCY\nDETECTION SYSTEM USING DEEP LEARNING")

add_authors(doc, "Author Name¹, Author Name², Author Name³, Author Name⁴, Guide Name⁵")
add_affiliation(doc, "¹˒²˒³˒⁴Student, Department of Computer Science & Engineering")
add_affiliation(doc, "⁵Associate Professor, Department of Computer Science & Engineering")
add_affiliation(doc, "RVR & JC College of Engineering, Chowdavaram, Guntur, Andhra Pradesh, India")
add_affiliation(doc, "Email: author1@gmail.com, author2@gmail.com, author3@gmail.com, author4@gmail.com")

# ── Abstract ──
add_abstract_block(
    doc,
    "This paper presents VitaDetect, an AI-powered system for the automated detection of vitamin "
    "deficiencies from images of human body parts -specifically nails, tongue, and skin. The system "
    "employs Convolutional Neural Networks (CNNs) trained on curated dermatological image datasets to "
    "classify visual symptoms into specific vitamin deficiency categories. Three independent deep "
    "learning models are developed: a Nail model capable of identifying Iodine Deficiency and Vitamin D "
    "Deficiency; a Tongue model detecting Vitamin B12 Deficiency and Iron Deficiency; and a Skin model "
    "recognizing Vitamin D Deficiency and Vitamin A Deficiency. The system is architected as a modern "
    "full-stack web application with a FastAPI-based RESTful backend serving TensorFlow/Keras models and "
    "a React-based frontend built with Vite. Upon analysis, the system provides the detected deficiency, "
    "associated disease correlation, confidence scores with a visual breakdown, a personalized weekly "
    "diet plan, and food recommendations. The decoupled architecture enables independent scaling of "
    "the inference engine and the user interface. This work demonstrates the feasibility of using "
    "computer vision techniques for preliminary nutritional screening as an accessible, non-invasive "
    "health assessment tool.",
    "Vitamin Deficiency, Convolutional Neural Network (CNN), Deep Learning, Image Classification, "
    "FastAPI, React, TensorFlow, Nutritional Health, Computer Vision"
)

# ── 1. INTRODUCTION ──
add_section_heading(doc, "1", "INTRODUCTION")
add_body_text(doc,
    "Vitamin deficiencies affect billions of people worldwide and are a leading cause of preventable "
    "diseases, particularly in developing nations. Traditional diagnosis methods require clinical "
    "blood tests, which can be expensive, time-consuming, and inaccessible in remote areas. However, "
    "many vitamin deficiencies manifest visible physical symptoms on the human body -discoloration "
    "and deformation of nails, changes in tongue appearance, and various skin conditions -that can "
    "be captured and analyzed through images."
)
add_body_text(doc,
    "Recent advancements in deep learning, specifically Convolutional Neural Networks (CNNs), have "
    "shown remarkable success in medical image analysis tasks, from diabetic retinopathy detection "
    "to skin cancer classification. Inspired by these developments, this work proposes VitaDetect, "
    "a system that leverages CNN-based image classification to detect potential vitamin deficiencies "
    "from photographs of nails, tongue, and skin."
)
add_body_text(doc,
    "The system is designed with a clear separation of concerns: a FastAPI REST API backend handles "
    "image processing and model inference using TensorFlow/Keras, while a modern React-based frontend "
    "built with Vite provides an intuitive drag-and-drop user interface. The architecture is "
    "production-ready, supporting CORS for cross-origin communication, environment-based configuration, "
    "and lazy model loading at startup. Additionally, the system generates personalized weekly diet "
    "plans based on the detected deficiency, offering users actionable health recommendations."
)

# ── 2. LITERATURE SURVEY ──
add_section_heading(doc, "2", "LITERATURE SURVEY")
add_body_text(doc,
    "The application of deep learning in dermatological image analysis has gained significant momentum "
    "in recent years. Esteva et al. (2017) demonstrated that a CNN trained on clinical images could "
    "achieve dermatologist-level classification of skin cancer, establishing a strong precedent for "
    "using deep learning in visual medical diagnosis. Similarly, research in automated tongue diagnosis "
    "has explored the use of CNNs for Traditional Chinese Medicine (TCM) applications, where tongue "
    "color and texture are analyzed for systemic health assessment."
)
add_body_text(doc,
    "Transfer learning using pre-trained architectures such as ResNet, VGG, MobileNet, and EfficientNet "
    "has become the standard approach for medical image classification, particularly when labeled datasets "
    "are limited. These architectures, pre-trained on ImageNet, provide rich feature representations "
    "that can be fine-tuned on domain-specific medical datasets."
)
add_body_text(doc,
    "From the reviewed literature, the following key observations emerge:"
)
add_bullet(doc, "Visual symptoms of nutritional deficiencies in nails, tongue, and skin are well-documented in clinical literature and are amenable to image-based classification.")
add_bullet(doc, "CNN-based approaches consistently outperform traditional machine learning methods (SVM, Random Forest) for image classification tasks in dermatology.")
add_bullet(doc, "Transfer learning significantly improves model accuracy when working with limited medical image datasets (typically under 2,000 images).")
add_bullet(doc, "Most existing systems are monolithic or notebook-based; few offer a production-ready web application with a decoupled frontend and backend architecture.")
add_bullet(doc, "There is a gap in integrating predictive models with actionable health recommendations, such as personalized diet plans based on the detected deficiency.")

# ── 3. PROPOSED SYSTEM ──
add_section_heading(doc, "3", "PROPOSED SYSTEM")
add_body_text(doc,
    "The proposed system, VitaDetect, is a full-stack web application designed to detect vitamin "
    "deficiencies from images of three body parts: nails, tongue, and skin. The system follows a "
    "decoupled client-server architecture with three main layers: the Frontend (React + Vite), the "
    "Backend (FastAPI), and the Inference Engine (TensorFlow/Keras CNN models)."
)
add_body_text(doc,
    "The user interacts with the React frontend to upload an image and select a body part. The image "
    "is sent via an HTTP POST request to the FastAPI backend's /api/analyze endpoint. The backend "
    "preprocesses the image -resizing it to 224×224 pixels, normalizing pixel values to [0,1] -and "
    "passes it to the appropriate CNN model based on the selected body part. The model returns class "
    "probabilities, from which the system determines the predicted deficiency, the associated disease "
    "correlation, and per-class confidence scores. A diet planner service then generates a personalized "
    "weekly meal plan based on the detected deficiency."
)

# System architecture figure
add_figure(doc, os.path.join(IMAGES_DIR, "uml 1.png"),
           "Fig 1: System Architecture -Component Interaction Diagram", Inches(5.2))

add_body_text(doc,
    "The system provides the following capabilities:"
)
add_bullet(doc, "Nail Analysis: Classifies nail images into three categories -No Vitamin Deficiency, Iodine Deficiency, and Vitamin D Deficiency -correlating to conditions such as bluish nails and alopecia areata.")
add_bullet(doc, "Tongue Analysis: Classifies tongue images into Vitamin B12 Deficiency and Iron Deficiency, correlating to diabetes indicators and pale tongue conditions.")
add_bullet(doc, "Skin Analysis: Classifies skin images into Vitamin D Deficiency and Vitamin A Deficiency, correlating to conditions such as acne and carcinoma.")
add_bullet(doc, "Diet Recommendations: Generates a weekly diet plan with day-wise food suggestions and a summary of recommended foods based on the identified deficiency.")
add_bullet(doc, "Interactive Confidence Visualization: Displays a bar chart breakdown of per-class confidence scores for transparent and interpretable predictions.")

# ── 4. METHODOLOGY ──
add_section_heading(doc, "4", "METHODOLOGY")
add_body_text(doc,
    "The methodology of the system is organized into the following stages:"
)

# 4.1 Dataset
p = doc.add_paragraph()
p.space_before = Pt(8)
p.space_after = Pt(4)
run = p.add_run("4.1 Dataset Preparation")
run.bold = True
run.font.size = Pt(10)
run.font.name = 'Times New Roman'

add_body_text(doc,
    "The dataset is organized into three body-part categories (Nail, Tongue, Skin) with the following "
    "class structure. The total dataset comprises 1,434 images split across training (1,016 images), "
    "validation (173 images), and test (245 images) sets."
)

# Dataset table
table = doc.add_table(rows=8, cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['Body Part', 'Class / Condition', 'Deficiency Mapping', 'Training Images']
data = [
    ['Nail', 'No Disease', 'No Vitamin Deficiency', '170'],
    ['Nail', 'Bluish Nail', 'Iodine Deficiency', '58'],
    ['Nail', 'Alopecia Areata', 'Vitamin D Deficiency', '76'],
    ['Tongue', 'Diabetes', 'Vitamin B12 Deficiency', '59'],
    ['Tongue', 'Pale Tongue', 'Iron Deficiency', '54'],
    ['Skin', 'Acne', 'Vitamin D Deficiency', '300'],
    ['Skin', 'Carcinoma', 'Vitamin A Deficiency', '299'],
]
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(9)
            run.font.name = 'Times New Roman'
for row_idx, row_data in enumerate(data):
    for col_idx, val in enumerate(row_data):
        cell = table.rows[row_idx + 1].cells[col_idx]
        cell.text = val
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'

cap_t = doc.add_paragraph()
cap_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
cap_t.space_after = Pt(8)
run = cap_t.add_run("Table 1: Dataset Class Distribution")
run.font.size = Pt(9)
run.font.name = 'Times New Roman'
run.bold = True

# 4.2 Image Preprocessing
p = doc.add_paragraph()
p.space_before = Pt(8)
p.space_after = Pt(4)
run = p.add_run("4.2 Image Preprocessing Pipeline")
run.bold = True
run.font.size = Pt(10)
run.font.name = 'Times New Roman'

add_body_text(doc,
    "Each uploaded image undergoes a standardized preprocessing pipeline before being fed into the "
    "CNN model. The pipeline consists of: (1) Resizing the image to 224×224 pixels to match the "
    "expected input dimensions of the CNN architecture, (2) Color space conversion from BGR to RGB, "
    "(3) Pixel normalization by scaling values to the [0, 1] range through division by 255, and "
    "(4) Batch dimension expansion by adding a leading axis to produce a tensor of shape (1, 224, 224, 3) "
    "suitable for model inference."
)

# Preprocessing pipeline figure
add_figure(doc, os.path.join(IMAGES_DIR, "uml 2.png"),
           "Fig 2: Image Preprocessing and Inference Pipeline", Inches(3.0))

# 4.3 Model Architecture
p = doc.add_paragraph()
p.space_before = Pt(8)
p.space_after = Pt(4)
run = p.add_run("4.3 Model Architecture and Training")
run.bold = True
run.font.size = Pt(10)
run.font.name = 'Times New Roman'

add_body_text(doc,
    "Three independent CNN models are trained, each specialized for a specific body part. The models "
    "are built using TensorFlow/Keras and saved in the .keras format. The nail model was trained for "
    "30 epochs with the best checkpoint selected at epoch 24. The tongue model was trained for 50 "
    "epochs with the first epoch checkpoint used. The skin model was trained for 46 epochs. All "
    "models use a 224×224 input image size and are served through a centralized PredictorService "
    "class that loads models at application startup."
)

# Model summary table
table2 = doc.add_table(rows=4, cols=5)
table2.style = 'Table Grid'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
headers2 = ['Model', 'Body Part', 'Input Size', 'Output Classes', 'Training Epochs']
data2 = [
    ['Nail CNN', 'Nails', '224×224', '3 (No Deficiency, Iodine, Vit D)', '30'],
    ['Tongue CNN', 'Tongue', '224×224', '2 (Vit B12, Iron)', '50'],
    ['Skin CNN', 'Skin', '224×224', '2 (Vit D, Vit A)', '46'],
]
for i, h in enumerate(headers2):
    cell = table2.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(9)
            run.font.name = 'Times New Roman'
for row_idx, row_data in enumerate(data2):
    for col_idx, val in enumerate(row_data):
        cell = table2.rows[row_idx + 1].cells[col_idx]
        cell.text = val
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'

cap_t2 = doc.add_paragraph()
cap_t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
cap_t2.space_after = Pt(8)
run = cap_t2.add_run("Table 2: Model Architecture Summary")
run.font.size = Pt(9)
run.font.name = 'Times New Roman'
run.bold = True

# 4.4 Backend Architecture
p = doc.add_paragraph()
p.space_before = Pt(8)
p.space_after = Pt(4)
run = p.add_run("4.4 Backend API Architecture")
run.bold = True
run.font.size = Pt(10)
run.font.name = 'Times New Roman'

add_body_text(doc,
    "The backend is built using the FastAPI framework, chosen for its high performance, automatic "
    "API documentation generation, and native support for asynchronous operations. The application "
    "uses a lifespan context manager to load all three TensorFlow models into memory at startup, "
    "eliminating cold-start latency during inference. The configuration system uses Pydantic Settings "
    "with environment variable overrides, making the application deployable across different environments."
)
add_body_text(doc,
    "The primary API endpoint POST /api/analyze accepts a multipart form upload containing the image "
    "file and a body_part parameter. The endpoint validates the input, decodes the image using OpenCV, "
    "passes it through the preprocessing and inference pipeline, and returns a structured JSON response "
    "containing the detected disease, vitamin deficiency, confidence score, per-class score breakdown, "
    "a weekly diet plan, and food recommendations."
)

# 4.5 Frontend Architecture
p = doc.add_paragraph()
p.space_before = Pt(8)
p.space_after = Pt(4)
run = p.add_run("4.5 Frontend Architecture")
run.bold = True
run.font.size = Pt(10)
run.font.name = 'Times New Roman'

add_body_text(doc,
    "The frontend is a single-page application built with React and bundled using Vite for fast "
    "development and optimized production builds. The interface features a drag-and-drop image "
    "uploader, a body part selector, and a comprehensive results panel. The results panel displays "
    "the detected body part, disease, vitamin deficiency, and confidence score, along with an "
    "interactive bar chart showing the confidence breakdown across all classes. A weekly diet plan "
    "and food recommendation section provides actionable health guidance. Communication with the "
    "backend is handled through an Axios HTTP client."
)

# ── 5. RESULTS ──
add_section_heading(doc, "5", "RESULTS")
add_body_text(doc,
    "The proposed VitaDetect system was successfully designed, implemented, and tested. The three "
    "CNN models for nail, tongue, and skin analysis were trained on the curated dataset and deployed "
    "through the FastAPI backend. The system was evaluated through end-to-end testing using the "
    "web interface."
)

# Result screenshot - Homepage
add_figure(doc, os.path.join(IMAGES_DIR, "Screenshot 2026-04-03 231933.png"),
           "Fig 3: VitaDetect Home Page -Image Upload Interface", Inches(5.2))

add_body_text(doc,
    "Figure 3 shows the landing page of the VitaDetect application. The interface presents a clean, "
    "modern dark-themed design with a centered image upload area that supports both drag-and-drop "
    "and click-to-browse functionality."
)

# Result screenshot - Analysis
add_figure(doc, os.path.join(IMAGES_DIR, "Screenshot 2026-04-03 232158.png"),
           "Fig 4: Skin Analysis Results -Disease Detection and Deficiency Identification", Inches(5.2))

add_body_text(doc,
    "Figure 4 demonstrates the analysis results for a skin image. The system correctly identified "
    "the body part as Skin, detected the disease as Acne, and classified the vitamin deficiency as "
    "Vitamin D Deficiency with a confidence score of 56.7%. The split-panel layout displays the "
    "uploaded image on the left and the analysis results on the right."
)

# Result screenshot - Confidence
add_figure(doc, os.path.join(IMAGES_DIR, "Screenshot 2026-04-03 232331.png"),
           "Fig 5: Confidence Breakdown -Per-Class Probability Distribution", Inches(3.5))

add_body_text(doc,
    "Figure 5 shows the confidence breakdown bar chart, providing a transparent view of the model's "
    "prediction probabilities across all classes. In this case, Vitamin D Deficiency received "
    "approximately 56.7% confidence while Vitamin A Deficiency received approximately 43.3%."
)

# Result screenshot - Diet Plan
add_figure(doc, os.path.join(IMAGES_DIR, "Screenshot 2026-04-03 232347.png"),
           "Fig 6: Weekly Diet Plan and Food Recommendations", Inches(3.5))

add_body_text(doc,
    "Figure 6 displays the personalized weekly diet plan generated based on the detected Vitamin D "
    "Deficiency. The system recommends foods such as Sun Bath (sunlight exposure), Mushrooms, Cheese, "
    "and Berries, distributed across a seven-day meal plan. A downloadable diet plan feature is also "
    "available for the user's convenience."
)

# ── Summary of Results ──
add_body_text(doc, "The key results of the system evaluation are summarized below:")
add_bullet(doc, "The system successfully loaded and served all three CNN models (Nail, Tongue, Skin) through the FastAPI backend with minimal startup latency.")
add_bullet(doc, "The React frontend provided an intuitive and responsive user experience with real-time image upload, body part selection, and results visualization.")
add_bullet(doc, "The confidence breakdown chart offered interpretable predictions, enabling users to understand the model's decision-making process.")
add_bullet(doc, "The diet recommendation engine generated contextually appropriate food suggestions for each detected deficiency.")
add_bullet(doc, "The system demonstrated end-to-end functionality from image upload to deficiency detection and dietary recommendation within seconds.")

# ── 6. CONCLUSION ──
add_section_heading(doc, "6", "CONCLUSION")
add_body_text(doc,
    "This work successfully developed VitaDetect, an AI-powered vitamin deficiency detection system "
    "that uses Convolutional Neural Networks to analyze images of nails, tongue, and skin for "
    "nutritional deficiency indicators. The system demonstrates the practical feasibility of using "
    "computer vision techniques for preliminary, non-invasive nutritional health screening."
)
add_body_text(doc,
    "The decoupled architecture -with a FastAPI backend serving TensorFlow models and a React "
    "frontend providing an interactive user interface -ensures maintainability, scalability, and "
    "independent evolution of the inference engine and the presentation layer. The inclusion of a "
    "personalized diet planning module adds immediate practical value by offering users actionable "
    "health recommendations alongside diagnostic information."
)
add_body_text(doc,
    "Future work includes expanding the model coverage to additional body parts such as hair, lips, "
    "and gums for more comprehensive nutritional assessment. Integrating larger and more diverse "
    "clinical datasets, employing advanced architectures such as EfficientNet and Vision Transformers, "
    "and deploying the system as a mobile application are planned enhancements that would increase "
    "the system's clinical relevance and accessibility."
)

# ── REFERENCES ──
add_section_heading(doc, "", "REFERENCES")
references = [
    'A. Esteva et al., "Dermatologist-level classification of skin cancer with deep neural networks," Nature, vol. 542, pp. 115–118, 2017.',
    'K. He, X. Zhang, S. Ren, and J. Sun, "Deep Residual Learning for Image Recognition," IEEE CVPR, 2016.',
    'M. Sandler et al., "MobileNetV2: Inverted Residuals and Linear Bottlenecks," IEEE CVPR, 2018.',
    'M. Tan and Q. V. Le, "EfficientNet: Rethinking Model Scaling for Convolutional Neural Networks," ICML, 2019.',
    'S. Ramírez, "FastAPI -Modern, fast web framework for building APIs with Python," 2019. Available: https://fastapi.tiangolo.com/',
    'Meta Platforms, "React -A JavaScript library for building user interfaces," 2023. Available: https://react.dev/',
    'M. Abadi et al., "TensorFlow: A System for Large-Scale Machine Learning," USENIX OSDI, 2016.',
    'F. Chollet, "Keras: The Python Deep Learning library," 2015. Available: https://keras.io/',
    'J. Deng et al., "ImageNet: A large-scale hierarchical image database," IEEE CVPR, 2009.',
    'G. Bradski, "The OpenCV Library," Dr. Dobb\'s Journal of Software Tools, 2000.',
    'A. Dosovitskiy et al., "An Image is Worth 16x16 Words: Transformers for Image Recognition at Scale," ICLR, 2021.',
    'World Health Organization, "Micronutrient deficiencies -Vitamin and Mineral Nutrition Information System," WHO, 2023.',
]

for i, ref in enumerate(references):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.space_after = Pt(2)
    pf = p.paragraph_format
    pf.left_indent = Cm(0.75)
    pf.first_line_indent = Cm(-0.75)
    run = p.add_run(f"{i+1}. {ref}")
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'

# ── Save ──
doc.save(OUTPUT_PATH)
print(f"[OK] Paper saved to: {OUTPUT_PATH}")
