"""
Build the final project report DOCX using template.docx as the base,
populating it with content from the docs/ folder and images from images/.
"""

import os
import copy
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import lxml.etree as etree

BASE_DIR = r"d:\All project files\All project files\new project"
DOCS_DIR = os.path.join(BASE_DIR, "docs")
IMAGES_DIR = os.path.join(BASE_DIR, "images")
TEMPLATE_PATH = os.path.join(BASE_DIR, "template.docx")
OUTPUT_PATH = os.path.join(BASE_DIR, "Vitamin_Deficiency_Detection_Report.docx")

# Image map: section name -> image filename
IMAGE_MAP = {
    "use_case": "use case diagram.jpeg",
    "sequence": "sequence diagram.jpeg",
    "state_chart": "state chart diagram.jpeg",
    "deployment": "deployment diagram.jpeg",
    "analysis_results": "analysis results.png",
    "confidence_chart": "confidence chart.png",
    "skin_detection": "skin detection image.png",
    "weekly_diet": "weekly diet plan.png",
    "front_page": "front page.png",
}


def read_md(path):
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


def add_page_break(doc):
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_break(docx_break_type())
    return para


def docx_break_type():
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    return br


def add_pg_break(doc):
    """Add a clean page break paragraph."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def set_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return h


def add_body_text(doc, text):
    """Add a paragraph with Body Text style."""
    p = doc.add_paragraph(style="Body Text")
    p.text = text
    return p


def add_normal(doc, text):
    p = doc.add_paragraph(text, style="Normal")
    return p


def parse_and_add_md_content(doc, md_text, skip_h1=True):
    """
    Parse markdown text and add content to the document.
    Handles: headings (##, ###), bullet lists, bold text, code blocks, tables, normal paragraphs.
    """
    lines = md_text.split("\n")
    in_code = False
    code_lines = []
    in_table = False
    table_data = []

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # Code block toggle
        if line.startswith("```"):
            if in_code:
                # End code block - add as monospaced body text
                if code_lines:
                    code_text = "\n".join(code_lines)
                    # Skip mermaid blocks entirely (no rendering support)
                    if not code_lines[0].strip().startswith("mermaid"):
                        p = doc.add_paragraph(style="Normal")
                        p.paragraph_format.left_indent = Inches(0.4)
                        run = p.add_run(code_text)
                        run.font.name = "Courier New"
                        run.font.size = Pt(9)
                code_lines = []
                in_code = False
            else:
                lang = line[3:].strip()
                in_code = True
                code_lines = [lang] if lang else []
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # Markdown table detection
        if "|" in line and line.strip().startswith("|"):
            # Collect table rows
            table_data = []
            while i < len(lines) and "|" in lines[i] and lines[i].strip().startswith("|"):
                row_line = lines[i].strip()
                # Skip separator rows (|---|---|)
                if re.match(r"^\|[\s\-\|]+\|$", row_line):
                    i += 1
                    continue
                cells = [c.strip() for c in row_line.strip("|").split("|")]
                table_data.append(cells)
                i += 1
            if table_data:
                ncols = max(len(r) for r in table_data)
                tbl = doc.add_table(rows=len(table_data), cols=ncols)
                tbl.style = "TableGrid"
                for ri, row in enumerate(table_data):
                    for ci, cell_text in enumerate(row):
                        if ci < ncols:
                            cell = tbl.cell(ri, ci)
                            cell.text = cell_text
                            if ri == 0:
                                for run in cell.paragraphs[0].runs:
                                    run.bold = True
                doc.add_paragraph()  # space after table
            continue

        # Headings
        if line.startswith("# ") and skip_h1:
            i += 1
            continue
        elif line.startswith("#### "):
            set_heading(doc, line[5:], level=4)
        elif line.startswith("### "):
            set_heading(doc, line[4:], level=3)
        elif line.startswith("## "):
            set_heading(doc, line[3:], level=2)
        elif line.startswith("# "):
            set_heading(doc, line[2:], level=1)

        # Bullet list
        elif line.startswith("- ") or line.startswith("* "):
            content = line[2:]
            p = doc.add_paragraph(style="List Paragraph")
            p.style = doc.styles["List Paragraph"]
            p.paragraph_format.left_indent = Inches(0.4)
            _add_inline_formatted_run(p, content)

        # Numbered list
        elif re.match(r"^\d+\.\s", line):
            content = re.sub(r"^\d+\.\s", "", line)
            p = doc.add_paragraph(style="List Paragraph")
            p.paragraph_format.left_indent = Inches(0.4)
            _add_inline_formatted_run(p, content)

        # Sub-bullet (indented)
        elif line.startswith("  - ") or line.startswith("  * "):
            content = line.strip()[2:]
            p = doc.add_paragraph(style="List Paragraph")
            p.paragraph_format.left_indent = Inches(0.8)
            _add_inline_formatted_run(p, content)

        # Empty line
        elif line.strip() == "":
            pass  # skip blank lines

        # Normal paragraph
        else:
            p = doc.add_paragraph(style="Body Text")
            _add_inline_formatted_run(p, line)

        i += 1


def _add_inline_formatted_run(paragraph, text):
    """
    Add text to a paragraph, handling **bold** and *italic* inline markdown.
    """
    # Pattern: **bold**, *italic*, `code`
    pattern = r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)"
    parts = re.split(pattern, text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        else:
            paragraph.add_run(part)


def add_image_to_doc(doc, image_filename, caption="", width=Inches(5.5)):
    img_path = os.path.join(IMAGES_DIR, image_filename)
    if not os.path.exists(img_path):
        print(f"  [WARN] Image not found: {img_path}")
        return
    try:
        doc.add_picture(img_path, width=width)
        # Center the image
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            cap_para = doc.add_paragraph(caption, style="Normal")
            cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cap_para.runs:
                run.italic = True
                run.font.size = Pt(9)
    except Exception as e:
        print(f"  [ERROR] Could not add image {image_filename}: {e}")


def update_cover_page(doc):
    """Update the cover page text of the loaded template."""
    for i, para in enumerate(doc.paragraphs):
        # Replace project title
        if "HOME APPLIANCES CONTROL USING RASPBERRY PI" in para.text:
            for run in para.runs:
                if "HOME APPLIANCES CONTROL USING RASPBERRY PI" in run.text:
                    run.text = run.text.replace(
                        "HOME APPLIANCES CONTROL USING RASPBERRY PI",
                        "VITAMIN DEFICIENCY DETECTION USING DEEP LEARNING"
                    )

        # Replace student names - simplified approach: replace known placeholder text
        if "S. R. PRABHANJAN" in para.text and "XXXXXXXXX" in para.text:
            for run in para.runs:
                run.text = ""
            para.runs[0].text = "Student Name\t\tRoll No."

        if para.text.strip().startswith("BEJJANKI VASHISTA"):
            for run in para.runs:
                run.text = ""

        if para.text.strip().startswith("JAHNAVI PATEL"):
            for run in para.runs:
                run.text = ""

        if para.text.strip().startswith("SATHYA HARSHA"):
            for run in para.runs:
                run.text = ""

        # Update guide name
        if "Ms. XXXX" in para.text:
            for run in para.runs:
                if "Ms. XXXX" in run.text:
                    run.text = run.text.replace("Ms. XXXX", "Ms. [Guide Name]")

        # Update declaration text
        if "Home a" in para.text and "dissertation" in para.text:
            for run in para.runs:
                run.text = ""
            para.runs[0].text = (
                "We hereby declare that the results embodied in the dissertation entitled "
                "\"Vitamin Deficiency Detection Using Deep Learning\" have been carried out "
                "by us under the guidance of Ms. [Guide Name], Assistant Professor, "
                "Department of Computer Science and Engineering. This work has not been "
                "submitted to any other University or Institution for the award of any "
                "degree or diploma."
            )

        # Update abstract
        if "Home Appliances is demonstrat" in para.text:
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = (
                    "This project presents a Vitamin Deficiency Detection System powered by "
                    "targeted Convolutional Neural Networks (CNNs). The system analyzes "
                    "user-submitted images of specific body parts -Nails, Tongue, or Skin -"
                    "to identify underlying vitamin and nutritional deficiencies. Users explicitly "
                    "select the body part category, which routes the image to the appropriate "
                    "specialized model, ensuring high accuracy while maintaining memory efficiency. "
                    "The system employs a modern decoupled architecture: a FastAPI backend for "
                    "REST API inference and a React (Vite) frontend for an intuitive user experience. "
                    "This approach eliminates the need for large language models (LLMs), drastically "
                    "reducing computational requirements without sacrificing diagnostic precision."
                )


def build_document():
    print("Loading template...")
    doc = Document(TEMPLATE_PATH)

    print("Updating cover page and front matter...")
    update_cover_page(doc)

    # Find the paragraph index of 'INTRODUCTION' heading to truncate from there
    intro_idx = None
    for i, para in enumerate(doc.paragraphs):
        if para.style.name == "Heading 1" and para.text.strip() == "INTRODUCTION":
            intro_idx = i
            break

    if intro_idx is None:
        print("[WARN] Could not find INTRODUCTION heading in template, appending to end.")
        intro_idx = len(doc.paragraphs) - 1

    # Remove all paragraphs from INTRODUCTION heading onwards
    # We need to work with the XML directly
    body = doc.element.body

    # Collect all block-level children (paragraphs and tables)
    body_children = list(body)

    # Find the XML element for our target paragraph
    target_para_elem = doc.paragraphs[intro_idx]._element

    # Find position in body_children
    target_pos = None
    for i, child in enumerate(body_children):
        if child is target_para_elem:
            target_pos = i
            break

    if target_pos is None:
        # Try finding the CHAPTER-1 heading
        for i, para in enumerate(doc.paragraphs):
            if "CHAPTER-1" in para.text or "CHAPTER 1" in para.text:
                target_para_elem = doc.paragraphs[i]._element
                for j, child in enumerate(body_children):
                    if child is target_para_elem:
                        target_pos = j
                        break
                break

    if target_pos is not None:
        print(f"Removing content from position {target_pos} onwards...")
        # Remove all elements from target_pos to end (keep sectPr last)
        elements_to_remove = []
        for child in body_children[target_pos:]:
            if child.tag.endswith("}sectPr"):
                continue  # keep last section properties
            elements_to_remove.append(child)
        for elem in elements_to_remove:
            body.remove(elem)
    else:
        print("[WARN] Could not find truncation point, content will be appended.")

    print("Adding Chapter 1: Introduction...")
    add_pg_break(doc)
    set_heading(doc, "CHAPTER - 1", level=1)
    set_heading(doc, "INTRODUCTION", level=1)

    # 1.1 Purpose
    set_heading(doc, "1.1 Purpose of the Project", level=2)
    parse_and_add_md_content(doc, read_md(os.path.join(DOCS_DIR, "1_Introduction", "1.1_Purpose_of_the_project.md")))

    # 1.2 Problem
    set_heading(doc, "1.2 Problem with Existing Systems", level=2)
    parse_and_add_md_content(doc, read_md(os.path.join(DOCS_DIR, "1_Introduction", "1.2_Problem_with_Existing_Systems.md")))

    # 1.3 Proposed System
    set_heading(doc, "1.3 Proposed System", level=2)
    parse_and_add_md_content(doc, read_md(os.path.join(DOCS_DIR, "1_Introduction", "1.3_Proposed_System.md")))

    # 1.4 Scope
    set_heading(doc, "1.4 Scope of the Project", level=2)
    parse_and_add_md_content(doc, read_md(os.path.join(DOCS_DIR, "1_Introduction", "1.4_Scope_of_the_Project.md")))

    # 1.5 Architecture Diagram
    set_heading(doc, "1.5 Architecture Diagram", level=2)
    p = doc.add_paragraph(
        "The system employs a modern decoupled client-server architecture where "
        "the React frontend communicates with the FastAPI backend via REST/JSON, "
        "and the backend routes the image to one of three specialized ML models "
        "based on the user-selected body part category.",
        style="Body Text"
    )
    # Add front page image as architecture illustration
    doc.add_paragraph()
    add_image_to_doc(doc, IMAGE_MAP["front_page"], "Fig 1.1: System Frontend -Vitamin Deficiency Detection Application")

    print("Adding Chapter 2: Literature Survey...")
    add_pg_break(doc)
    set_heading(doc, "CHAPTER - 2", level=1)
    set_heading(doc, "LITERATURE SURVEY", level=1)
    set_heading(doc, "2.1 Literature Survey", level=2)
    parse_and_add_md_content(doc, read_md(os.path.join(DOCS_DIR, "2_Literature_Survey", "2.1_Literature_Survey.md")))

    print("Adding Chapter 3: SRS...")
    add_pg_break(doc)
    set_heading(doc, "CHAPTER - 3", level=1)
    set_heading(doc, "SOFTWARE REQUIREMENT SPECIFICATION", level=1)

    set_heading(doc, "3.1 Introduction to SRS", level=2)
    parse_and_add_md_content(doc, read_md(os.path.join(DOCS_DIR, "3_Software_Requirement_Specification", "3.1_Introduction_to_SRS.md")))

    set_heading(doc, "3.2 Role of SRS", level=2)
    parse_and_add_md_content(doc, read_md(os.path.join(DOCS_DIR, "3_Software_Requirement_Specification", "3.2_Role_of_SRS.md")))

    set_heading(doc, "3.3 Requirements Specification Document", level=2)
    parse_and_add_md_content(doc, read_md(os.path.join(DOCS_DIR, "3_Software_Requirement_Specification", "3.3_Requirements_Specification_Document.md")))

    set_heading(doc, "3.4 Functional Requirements", level=2)
    parse_and_add_md_content(doc, read_md(os.path.join(DOCS_DIR, "3_Software_Requirement_Specification", "3.4_Functional_Requirements.md")))

    set_heading(doc, "3.5 Non-Functional Requirements", level=2)
    parse_and_add_md_content(doc, read_md(os.path.join(DOCS_DIR, "3_Software_Requirement_Specification", "3.5_Non-Functional_Requirements.md")))

    set_heading(doc, "3.6 Performance Requirements", level=2)
    parse_and_add_md_content(doc, read_md(os.path.join(DOCS_DIR, "3_Software_Requirement_Specification", "3.6_Performance_Requirements.md")))

    set_heading(doc, "3.7 Software Requirements", level=2)
    parse_and_add_md_content(doc, read_md(os.path.join(DOCS_DIR, "3_Software_Requirement_Specification", "3.7_Software_Requirements.md")))

    set_heading(doc, "3.8 Hardware Requirements", level=2)
    parse_and_add_md_content(doc, read_md(os.path.join(DOCS_DIR, "3_Software_Requirement_Specification", "3.8_Hardware_Requirements.md")))

    print("Adding Chapter 4: System Design...")
    add_pg_break(doc)
    set_heading(doc, "CHAPTER - 4", level=1)
    set_heading(doc, "SYSTEM DESIGN", level=1)

    set_heading(doc, "4.1 Introduction to UML", level=2)
    parse_and_add_md_content(doc, read_md(os.path.join(DOCS_DIR, "4_System_Design", "4.1_Introduction_to_UML.md")))

    set_heading(doc, "4.2 UML Diagrams", level=2)
    doc.add_paragraph(
        "The following sections present the UML diagrams used to model the Vitamin Deficiency Detection System.",
        style="Body Text"
    )

    set_heading(doc, "4.3 Use Case Diagram", level=2)
    doc.add_paragraph(
        "The use case diagram illustrates the interaction between the User and System Admin with the system. "
        "The primary user flow begins with selecting a body part category, followed by uploading an image, "
        "triggering preprocessing and model routing, and finally viewing the results.",
        style="Body Text"
    )
    doc.add_paragraph()
    add_image_to_doc(doc, IMAGE_MAP["use_case"], "Fig 4.1: Use Case Diagram -Vitamin Deficiency Detection System")

    set_heading(doc, "4.4 Sequence Diagram", level=2)
    doc.add_paragraph(
        "The sequence diagram shows the interaction flow between the User, React Frontend, FastAPI Backend, "
        "and ML Models. The user's category selection and image upload are sent as a multipart POST request "
        "to /api/analyze, which preprocesses the image and routes it to the selected model.",
        style="Body Text"
    )
    doc.add_paragraph()
    add_image_to_doc(doc, IMAGE_MAP["sequence"], "Fig 4.2: Sequence Diagram -User to Backend ML Model Interaction")

    set_heading(doc, "4.5 State Chart Diagram", level=2)
    doc.add_paragraph(
        "The state chart diagram captures the different states of the application: Idle → Selecting → "
        "Uploading → Processing → Analyzing → Success/Error → Idle.",
        style="Body Text"
    )
    doc.add_paragraph()
    add_image_to_doc(doc, IMAGE_MAP["state_chart"], "Fig 4.3: State Chart Diagram -Application State Transitions")

    set_heading(doc, "4.6 Deployment Diagram", level=2)
    doc.add_paragraph(
        "The deployment diagram shows the hosting configuration: a client browser communicates with "
        "a Vite/React static file server and a Uvicorn/FastAPI backend that loads TensorFlow models.",
        style="Body Text"
    )
    doc.add_paragraph()
    add_image_to_doc(doc, IMAGE_MAP["deployment"], "Fig 4.4: Deployment Diagram -System Hosting Architecture")

    set_heading(doc, "4.7 Technologies Used", level=2)
    parse_and_add_md_content(doc, read_md(os.path.join(DOCS_DIR, "4_System_Design", "4.7_Technologies_Used.md")))

    print("Adding Chapter 5: Implementation...")
    add_pg_break(doc)
    set_heading(doc, "CHAPTER - 5", level=1)
    set_heading(doc, "IMPLEMENTATION", level=1)

    set_heading(doc, "5.1 Setup and Connections", level=2)
    parse_and_add_md_content(doc, read_md(os.path.join(DOCS_DIR, "5_Implementation", "5.1_Setup_and_Connections.md")))

    set_heading(doc, "5.2 Coding the Logic", level=2)
    parse_and_add_md_content(doc, read_md(os.path.join(DOCS_DIR, "5_Implementation", "5.2_Coding_the_logic.md")))

    set_heading(doc, "5.3 Connecting the Frontend", level=2)
    parse_and_add_md_content(doc, read_md(os.path.join(DOCS_DIR, "5_Implementation", "5.3_Connecting_the_frontend.md")))

    set_heading(doc, "5.4 Screenshots", level=2)
    doc.add_paragraph(
        "The following screenshots demonstrate the key interfaces of the Vitamin Deficiency Detection System.",
        style="Body Text"
    )
    doc.add_paragraph()
    add_image_to_doc(doc, IMAGE_MAP["skin_detection"], "Fig 5.1: Skin Deficiency Detection -Analysis Interface")
    doc.add_paragraph()
    add_image_to_doc(doc, IMAGE_MAP["analysis_results"], "Fig 5.2: Analysis Results -Deficiency Detection Output")
    doc.add_paragraph()
    add_image_to_doc(doc, IMAGE_MAP["confidence_chart"], "Fig 5.3: Confidence Chart -Model Prediction Probabilities")
    doc.add_paragraph()
    add_image_to_doc(doc, IMAGE_MAP["weekly_diet"], "Fig 5.4: Weekly Diet Plan -Nutritional Recommendation Output")

    print("Adding Chapter 6: Software Testing...")
    add_pg_break(doc)
    set_heading(doc, "CHAPTER - 6", level=1)
    set_heading(doc, "SOFTWARE TESTING", level=1)

    set_heading(doc, "6.1 Introduction to Testing", level=2)
    parse_and_add_md_content(doc, read_md(os.path.join(DOCS_DIR, "6_Software_Testing", "6.1_Introduction.md")))

    set_heading(doc, "6.2 Testing Objectives", level=2)
    parse_and_add_md_content(doc, read_md(os.path.join(DOCS_DIR, "6_Software_Testing", "6.2_Testing_Objectives.md")))

    set_heading(doc, "6.3 Testing Strategies", level=2)
    parse_and_add_md_content(doc, read_md(os.path.join(DOCS_DIR, "6_Software_Testing", "6.3_Testing_Strategies.md")))

    set_heading(doc, "6.4 System Evaluation", level=2)
    parse_and_add_md_content(doc, read_md(os.path.join(DOCS_DIR, "6_Software_Testing", "6.4_System_Evaluation.md")))

    set_heading(doc, "6.5 Testing New System", level=2)
    parse_and_add_md_content(doc, read_md(os.path.join(DOCS_DIR, "6_Software_Testing", "6.5_Testing_New_System.md")))

    set_heading(doc, "6.6 Test Cases", level=2)
    parse_and_add_md_content(doc, read_md(os.path.join(DOCS_DIR, "6_Software_Testing", "6.6_Test_Cases.md")))

    print("Adding Chapter 7: Conclusion...")
    add_pg_break(doc)
    set_heading(doc, "CONCLUSION", level=1)
    parse_and_add_md_content(doc, read_md(os.path.join(DOCS_DIR, "7_Conclusion", "7.1_Conclusion.md")))

    print("Adding Chapter 8: Future Enhancements...")
    add_pg_break(doc)
    set_heading(doc, "FUTURE ENHANCEMENTS", level=1)
    parse_and_add_md_content(doc, read_md(os.path.join(DOCS_DIR, "8_Future_Enhancements", "8.1_Future_Enhancements.md")))

    print("Adding References...")
    add_pg_break(doc)
    set_heading(doc, "REFERENCES", level=1)
    parse_and_add_md_content(doc, read_md(os.path.join(DOCS_DIR, "9_References", "9.1_References.md")))

    print("Adding Bibliography...")
    add_pg_break(doc)
    set_heading(doc, "BIBLIOGRAPHY", level=1)
    parse_and_add_md_content(doc, read_md(os.path.join(DOCS_DIR, "10_Bibliography", "10.1_Bibliography.md")))

    print(f"\nSaving document to: {OUTPUT_PATH}")
    doc.save(OUTPUT_PATH)
    print("Done! Document saved successfully.")
    print(f"Output: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_document()
